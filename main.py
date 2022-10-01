# coding=utf-8
import time
import docx
from docx.shared import Cm
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By

options = webdriver.ChromeOptions()
options.headless = True  # 關閉顯示
options.add_experimental_option("excludeSwitches", ["enable-logging"])
options.add_experimental_option("prefs", {"profile.default_content_setting_values.notifications": 2})
options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/104.0.5112.81 Safari/537.36") # https://www.whatismybrowser.com/detect/what-is-my-user-agent/
options.add_argument("window-size=1920,1080")
service = Service(r"C:/chromedriver_win32/chromedriver.exe")
driver = webdriver.Chrome(service=service, options=options)

# FB 帳號密碼
FB_ACCOUNT = "YOUR_FB_ACCOUNT"
FB_PASSWORD = "YOUR_FB_PASSWORD"

# 輸入檔案
INPUT_FILE = "pages.txt"

LOGIN_URL = "https://www.facebook.com/login.php"
POSTS_CONTAINER_XPATH = "//div[@role='feed']" # 文章容器
POSTS_XPATH = "(//div[@role='main'])[2]/div/div" # 文章另類選擇
GROUP_POSTS_XPATH = "(//div[@role='main'])[2]/div/div" # 文章另類選擇

CHAT_XPATH = "//div[@data-testid='mwchat-tabs']"
SHOW_MORE_ZH_XPATH = "//div[@role='button'][contains(text(),'顯示更多')]"   # 顯示更多
SHOW_MORE_EN_XPATH = "//div[@role='button'][normalize-space()='See more']" # 顯示更多
ALBUM_BANNER_ZH_XPATH = "//span[contains(text(),'Album')]" # 相簿條
ALBUM_BANNER_EN_XPATH = "//span[contains(text(),'相簿')]" # 相簿條
COMMENTS_CONTAINER_XPATH = "//div[@data-visualcompletion='ignore-dynamic']" # 讚/留言
BANNER_XPATH = "//div[contains(@style,'z-index: 1;')]" # 橫條
IMAGES = []
URLS = []
S = lambda X: driver.execute_script("return document.body.parentNode.scroll" + X)


class Facebook:
    def __init__(self, email: str, password: str) -> None:
        """
        初始化

        Args:
            email (str): FB帳號Email
            password (str): FB密碼
        """
        self.email = email
        self.password = password

        driver.get(LOGIN_URL)
        time.sleep(1)

    def login(self) -> None:
        """
        登入FB
        """
        print("Login...")
        driver.find_element(By.ID, "email").send_keys(self.email)  # Give keyboard input
        driver.find_element(By.ID, "pass").send_keys(self.password)  # Give password as input too

        login_button = driver.find_element(By.ID, "loginbutton")
        login_button.click()
        time.sleep(2)

    def doScrolling(self) -> None:
        """
        滑動畫面
        """
        driver.execute_script("window.scrollTo(0,document.body.scrollHeight)")
        time.sleep(2)
        print("Scrolling...")

    def getGroupScreenShots(self, url: str) -> None:
        """
        擷取文章圖片

        Args:
            url (str): 要擷取文章的網址
        """
        driver.get(f"{url}?sorting_setting=CHRONOLOGICAL")
        time.sleep(2)

        # 載入至少三篇文章
        while True:
            self.doScrolling()
            time.sleep(2)
            try:
                # print("way 1")
                all_post_container = driver.find_elements(By.XPATH, POSTS_CONTAINER_XPATH)
                all_post_container = all_post_container[0] if len(all_post_container) == 1 else all_post_container[1]
                posts = all_post_container.find_elements(By.XPATH, './div')[1:-2]
            except Exception:
                # print("way 2")
                posts = driver.find_elements(By.XPATH, POSTS_XPATH)[:-4]
            post_count = len(posts)
            print(f"Load {post_count} posts.")
            if post_count >= 3:
                break

        self.removeAlbum()
        self.removeCommentAndLike()
        self.removeBanner()
        self.clickShowMore()
        self.removeChatRoom()

        # 拍取文章照片
        name = url.rsplit("/", 1)[-1]
        for i in range(3 if post_count > 3 else post_count):
            time.sleep(2)
            driver.execute_script("arguments[0].scrollIntoView();", posts[i])
            path = f"output/{name}_{i}.png"
            print(f"Saving screen at {path}")
            posts[i].screenshot(path)
            IMAGES.append(path)
            URLS.append(url)

    def removeAlbum(self) -> None:
        """移除相簿橫條
        """
        albums = driver.find_elements(By.XPATH, ALBUM_BANNER_ZH_XPATH)
        for album in albums:
            driver.execute_script("""var element = arguments[0];element.parentNode.removeChild(element);""", album)

        albums = driver.find_elements(By.XPATH, ALBUM_BANNER_EN_XPATH)
        for album in albums:
            driver.execute_script("""var element = arguments[0];element.parentNode.removeChild(element);""", album)

    def removeCommentAndLike(self) -> None:
        """移除點讚留言
        """
        comments = driver.find_elements(By.XPATH, COMMENTS_CONTAINER_XPATH)
        for comment in comments:
            driver.execute_script("""var element = arguments[0];element.parentNode.removeChild(element);""", comment)

    def removeBanner(self) -> None:
        """移除 Banner
        """
        driver.execute_script("document.querySelector('div[role=\"banner\"]').remove()")
        banners = driver.find_elements(By.XPATH, BANNER_XPATH)
        for banner in banners:
            driver.execute_script("""var element = arguments[0];element.parentNode.removeChild(element);""", banner)

    def removeChatRoom(self) -> None:
        """關閉聊天室
        """
        time.sleep(2)
        chat_elem = driver.find_elements(By.XPATH, CHAT_XPATH)
        for e in chat_elem:
            driver.execute_script("""var element = arguments[0];element.parentNode.removeChild(element);""", e)

    def clickShowMore(self) -> None:
        """點擊'觀看更多'
        """
        show_more_elem = driver.find_elements(By.XPATH, SHOW_MORE_EN_XPATH)
        for e in show_more_elem:
            try:
                e.click()
            except Exception:
                pass
        show_more_elem = driver.find_elements(By.XPATH, SHOW_MORE_ZH_XPATH)
        for e in show_more_elem:
            try:
                e.click()
            except Exception:
                pass

    def save2docx(self) -> None:
        """保存圖片到WORD
        """
        doc = docx.Document()
        for i in range(len(IMAGES)):
            doc.add_heading(f"{URLS[i]} {i%3}", level=2)
            doc.add_picture(IMAGES[i], width=Cm(18))
        fn = datetime.now().strftime("%Y_%m_%d-%I_%M_%S_%p")
        doc.save(f"{fn}.docx")


    def quit(self) -> None:
        """
        關閉驅動
        """
        driver.quit()


if __name__ == "__main__":
    fb = Facebook(email=FB_ACCOUNT, password=FB_PASSWORD)
    fb.login()
    page_urls = open(INPUT_FILE, encoding="utf-8").read().splitlines()
    for page_url in page_urls:
        fb.getGroupScreenShots(page_url)
    fb.save2docx()
    fb.quit()
